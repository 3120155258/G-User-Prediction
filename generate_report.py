"""
生成作业分析报告 Word 文档
"""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), 'libs'))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置默认段落间距
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading

def add_para(text, bold=False, size=None, color=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1)
    for run in p.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def set_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 封面 ==========
for _ in range(6):
    doc.add_paragraph()

add_para('机器学习小组作业', bold=True, size=26, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('——分析报告', bold=True, size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('题目：5G 用户预测', bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para(f'软件工程系', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f'文档生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ========== 目录 ==========
add_heading_styled('目  录', level=1)
toc_items = [
    '一、作业概述（是什么）',
    '    1.1 作业基本信息',
    '    1.2 问题背景',
    '    1.3 任务目标',
    '    1.4 核心要求',
    '二、作业目的与意义（为什么）',
    '    2.1 教学目标',
    '    2.2 实际应用价值',
    '三、实施指南（怎么做）',
    '    3.1 数据说明',
    '    3.2 评估指标',
    '    3.3 推荐技术路线',
    '    3.4 报告撰写要求',
    '四、提交要求与评分',
    '    4.1 提交清单',
    '    4.2 评分与加分项',
    '    4.3 重要时间节点',
    '五、关键信息汇总表',
]
for item in toc_items:
    add_para(item, size=11)

doc.add_page_break()

# ========== 第一章：作业概述 ==========
add_heading_styled('一、作业概述（是什么）', level=1)

add_heading_styled('1.1 作业基本信息', level=2)
info_table = doc.add_table(rows=6, cols=2, style='Table Grid')
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('项目', '内容'),
    ('作业题目', '5G 用户预测'),
    ('课程名称', '人工智能'),
    ('所属院系', '软件工程系'),
    ('提交日期', '第16周（6月15日）随堂提交'),
    ('组织形式', '小组完成（建议5~6人/组，重修同学可1人1组）'),
]
for i, (key, val) in enumerate(info_data):
    info_table.rows[i].cells[0].text = key
    info_table.rows[i].cells[1].text = val
    if i == 0:
        set_cell_shading(info_table.rows[i].cells[0], '4472C4')
        set_cell_shading(info_table.rows[i].cells[1], '4472C4')
        for cell in info_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.bold = True
set_table_font(info_table)
doc.add_paragraph()

add_heading_styled('1.2 问题背景', level=2)
add_para('在2022世界互联网大会乌镇峰会期间发布的《世界互联网发展报告2022》显示，2022年第一季度，全球5G用户数增加7000万人，总数达到6.2亿人左右，5G人口覆盖率超25%。据爱立信及全球移动通信系统协会（GSMA）预测，到2022年底，全球5G用户数量将突破10亿人。')
add_para('对于通信运营商来说，面对如此庞大的5G市场，如何基于用户侧信息进行用户画像，再进一步对潜在5G使用者进行精准营销十分有益。本作业即围绕这一实际商业场景展开。')

add_heading_styled('1.3 任务目标', level=2)
add_para('根据用户基本信息和通信相关数据（包括用户话费信息、流量使用、活跃行为、套餐类型、区域信息等特征字段），通过训练数据训练机器学习模型，预测测试集中每个样本是否属于5G用户。', bold=True)
add_para('本质上是一个二分类问题（5G用户/非5G用户），也可转化为0~1回归问题。')

add_heading_styled('1.4 核心要求', level=2)

add_para('硬性要求（必须完成）：', bold=True)
add_bullet('实现不少于两种人工智能算法，对5G用户进行预测')
add_bullet('撰写相应分析报告（双面打印，不多于6页）')
add_bullet('源码压缩包至少包含一份 Jupyter Notebook 文件，展示大致建模过程和实验结果')
add_bullet('其他自定义文件以普通 Python (.py) 文件包含在压缩包内')
add_bullet('报告纸质版于第16周随堂提交')

doc.add_paragraph()
add_para('提示与建议：', bold=True)
add_bullet('用户预测问题可抽象为二分类问题或0~1的回归问题')
add_bullet('报告中可包含：任务思考、数据分析、模型选择原因、查阅资料的总结、不同模型比较、结果分析、改进思路等')
add_bullet('报告中不能包含实验过程源码，如有需要可包含简洁的伪代码')
add_bullet('撰写报告时建议图文并茂，适当加入图表')
add_bullet('可使用未学过的AI模型或算法，但必须在报告中给出模型原理、选择原因、结果分析')
add_bullet('可使用任意实验课中未涉及的Python AI库，不必给出详细说明')

doc.add_page_break()

# ========== 第二章：目的与意义 ==========
add_heading_styled('二、作业目的与意义（为什么）', level=1)

add_heading_styled('2.1 教学目标', level=2)
add_para('本作业的核心教学目的：', bold=True)
goals = [
    ('巩固理论知识', '通过Python编程环境实现AI解决方案，将课堂所学的机器学习理论应用于实际问题'),
    ('熟悉数据操作', '掌握对数据集的基本操作，包括数据加载、清洗、特征工程等完整数据处理流程'),
    ('掌握AI工具', '熟悉相关AI Python库（如scikit-learn、XGBoost、LightGBM等）的使用'),
    ('培养分析思维', '对实际问题中数据特性与目标的思考，学会如何选择合适模型和评估方法'),
    ('模型比较能力', '不同模型间的比较和分析，理解各算法的优劣和适用场景'),
    ('团队协作能力', '分组完成，培养团队分工协作、项目管理的能力'),
]
for title, desc in goals:
    p = doc.add_paragraph()
    run_title = p.add_run(f'● {title}：')
    run_title.bold = True
    run_title.font.name = '宋体'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_desc = p.add_run(desc)
    run_desc.font.name = '宋体'
    run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading_styled('2.2 实际应用价值', level=2)
add_para('本作业模拟了通信运营商的实际业务场景：')
add_bullet('用户画像构建：基于多维度用户数据构建精准的用户画像，是数据驱动营销的基础能力')
add_bullet('二分类预测：5G用户预测本质上是一个典型的机器学习二分类问题，是业界最常见的AI应用场景之一')
add_bullet('AUC评估：实际业务中常使用AUC作为模型评估指标，因为它不受类别不平衡和阈值选择的影响')
add_bullet('工程实践：从数据预处理到模型部署的完整pipeline，体现了工业界ML项目的全流程')

doc.add_page_break()

# ========== 第三章：实施指南 ==========
add_heading_styled('三、实施指南（怎么做）', level=1)

add_heading_styled('3.1 数据说明', level=2)
add_para('数据集文件：train.csv，包含60个字段。', bold=True)

data_table = doc.add_table(rows=5, cols=3, style='Table Grid')
data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
data_info = [
    ('特征类别', '字段名', '字段描述'),
    ('样本标识', 'id', '样本标识ID'),
    ('离散型特征', 'cat_0 ~ cat_19', '共20个离散型（类别型）特征'),
    ('数值型特征', 'num_0 ~ num_37', '共38个数值型特征'),
    ('目标字段', 'target', '是否为5G用户（0或1）'),
]
for i, (a, b, c) in enumerate(data_info):
    data_table.rows[i].cells[0].text = a
    data_table.rows[i].cells[1].text = b
    data_table.rows[i].cells[2].text = c
    if i == 0:
        for cell in data_table.rows[i].cells:
            set_cell_shading(cell, '4472C4')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.bold = True
set_table_font(data_table)

doc.add_paragraph()
add_para('特征字段涵盖内容：用户话费信息、流量使用情况、活跃行为、套餐类型、区域信息等。')

add_heading_styled('3.2 评估指标', level=2)
add_para('本次作业的评价标准采用 AUC (Area Under ROC Curve)，分数越高，效果越好。', bold=True)
add_para('AUC（ROC曲线下面积）是二分类模型最常用的评估指标之一，取值范围为[0,1]。其优势在于：')
add_bullet('不受分类阈值选择的影响')
add_bullet('对类别不平衡数据具有较好的鲁棒性')
add_bullet('直观反映模型区分正负样本的排序能力')

doc.add_paragraph()
add_para('评估代码参考：', bold=True)
p = doc.add_paragraph()
run = p.add_run('''from sklearn.metrics import roc_auc_score
auc = roc_auc_score(y_true, y_pred_proba)''')
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_heading_styled('3.3 推荐技术路线', level=2)
add_para('阶段一：数据探索与预处理', bold=True, size=12)
add_bullet('加载数据，查看基本统计信息（样本量、特征维度、数据类型）')
add_bullet('分析目标变量分布，评估类别平衡情况')
add_bullet('检查缺失值、异常值，进行数据清洗')
add_bullet('离散特征编码（Label Encoding / One-Hot Encoding）')
add_bullet('数值特征标准化/归一化（可选）')
add_bullet('特征相关性分析，筛选高价值特征')

doc.add_paragraph()
add_para('阶段二：模型训练', bold=True, size=12)
add_bullet('选择至少2种AI算法（如：Logistic Regression、Random Forest、XGBoost、LightGBM、神经网络等）')
add_bullet('划分训练集/验证集（建议使用分层采样 stratified split）')
add_bullet('训练各模型并调参')

doc.add_paragraph()
add_para('阶段三：模型评估与比较', bold=True, size=12)
add_bullet('使用AUC作为主要评估指标')
add_bullet('绘制ROC曲线对比各模型表现')
add_bullet('交叉验证（K-Fold CV）评估模型稳定性')
add_bullet('分析特征重要性，理解哪些用户特征对5G使用预测影响最大')

doc.add_paragraph()
add_para('阶段四：报告撰写', bold=True, size=12)
add_bullet('整理实验过程与分析结果')
add_bullet('撰写结构化报告（不包含源码，可使用伪代码）')
add_bullet('图文并茂，包含关键图表')

add_heading_styled('3.4 报告撰写要求', level=2)
report_items = [
    '对任务的思考和理解',
    '数据的详细分析（分布、相关性、缺失值等）',
    '模型选择的原因和理论依据',
    '对查阅的相关资料的思考和总结',
    '不同模型的性能比较和结果分析',
    '可改进思路的思考（未来工作）',
    '图文并茂，适当加入图表',
    '不包含实验源码（可包含简洁伪代码）',
    '双面打印，总页数不超过6页',
]
for item in report_items:
    add_bullet(item)

doc.add_page_break()

# ========== 第四章：提交与评分 ==========
add_heading_styled('四、提交要求与评分', level=1)

add_heading_styled('4.1 提交清单', level=2)

submit_table = doc.add_table(rows=5, cols=3, style='Table Grid')
submit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
submit_data = [
    ('提交项', '形式', '说明'),
    ('源码压缩包', '电子版', '至少含1份Jupyter Notebook文件 + Python .py文件，由班长/学委汇总发送至 ruihong@fjnu.edu.cn'),
    ('电子版报告', '电子版', '包含在源码压缩包中'),
    ('纸质版报告', '纸质', '第16周随堂提交，双面打印，不超过6页'),
    ('成员分工与小组自评表', '电子版', '每小组1份，与其他材料一起提交'),
]
for i, (a, b, c) in enumerate(submit_data):
    submit_table.rows[i].cells[0].text = a
    submit_table.rows[i].cells[1].text = b
    submit_table.rows[i].cells[2].text = c
    if i == 0:
        for cell in submit_table.rows[i].cells:
            set_cell_shading(cell, '4472C4')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.bold = True
set_table_font(submit_table)

doc.add_paragraph()

add_para('分组名单提交：', bold=True)
add_para('组长于4月19日20时59分前将分组名单提交给班长或学习委员，汇总后发送至邮箱 ruihong@fjnu.edu.cn。')

add_heading_styled('4.2 评分与加分项', level=2)

add_para('评分标准：', bold=True)
add_bullet('模型实现质量：是否实现≥2种AI算法，模型选择是否合理')
add_bullet('代码质量：结构清晰、注释规范、逻辑严谨')
add_bullet('报告质量：分析深度、图文质量、论述逻辑')
add_bullet('AUC指标表现：模型在验证集上的预测性能')

doc.add_paragraph()
add_para('加分项（最多+5分）：', bold=True, color=(192, 0, 0))
add_bullet('小组作业上传GitHub开源')
add_bullet('撰写相关英文 ReadMe 文件')
add_bullet('进行必要英文注释')
add_bullet('注：依据组员分工不同，小组成员所加分数会有差异（参考成员分工与小组自评表）')
add_bullet('开源地址需在实验报告中标明')

add_heading_styled('4.3 重要时间节点', level=2)

timeline_table = doc.add_table(rows=5, cols=3, style='Table Grid')
timeline_table.alignment = WD_TABLE_ALIGNMENT.CENTER
timeline_data = [
    ('时间节点', '事项', '说明'),
    ('4月19日 20:59前', '提交分组名单', '组长提交给班长/学委，汇总后发 ruihong@fjnu.edu.cn'),
    ('第16周（6月15日）', '随堂提交', '源码压缩包 + 纸质报告 + 自评表'),
    ('第16或17周实验课', 'PPT汇报', '组长汇报展示，每组10~15分钟'),
    ('截止时间后', '缺交处理', '超过截止时间的附上缺交名单'),
]
for i, (a, b, c) in enumerate(timeline_data):
    timeline_table.rows[i].cells[0].text = a
    timeline_table.rows[i].cells[1].text = b
    timeline_table.rows[i].cells[2].text = c
    if i == 0:
        for cell in timeline_table.rows[i].cells:
            set_cell_shading(cell, '4472C4')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.bold = True
set_table_font(timeline_table)

doc.add_page_break()

# ========== 第五章：关键信息汇总 ==========
add_heading_styled('五、关键信息汇总表', level=1)

summary_table = doc.add_table(rows=25, cols=2, style='Table Grid')
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 调整列宽
for row in summary_table.rows:
    row.cells[0].width = Cm(4.5)
    row.cells[1].width = Cm(11)

summary_data = [
    ('维度', '详细内容'),
    ('', ''),
    ('【基本信息】', ''),
    ('作业题目', '5G 用户预测'),
    ('课程', '人工智能'),
    ('院系', '软件工程系'),
    ('问题类型', '二分类问题（预测用户是否为5G用户）'),
    ('', ''),
    ('【数据信息】', ''),
    ('数据文件', 'train.csv'),
    ('特征数量', '60个字段（id + 20个离散型cat_ + 38个数值型num_ + target）'),
    ('目标字段', 'target（0=非5G用户，1=5G用户）'),
    ('特征内容', '话费信息、流量使用、活跃行为、套餐类型、区域信息等'),
    ('', ''),
    ('【技术要求】', ''),
    ('编程语言', 'Python'),
    ('最少算法数', '不少于2种人工智能算法'),
    ('评估指标', 'AUC (ROC曲线下面积)'),
    ('核心交付物', 'Jupyter Notebook + Python文件 + 分析报告'),
    ('', ''),
    ('【提交要求】', ''),
    ('分组', '建议5~6人/组（重修可1人1组）'),
    ('提交日期', '第16周（6月15日）随堂提交'),
    ('报告格式', '双面打印，不超过6页，不含源码'),
    ('加分项', 'GitHub开源 + 英文README（最多+5分）'),
]

for i, (key, val) in enumerate(summary_data):
    summary_table.rows[i].cells[0].text = key
    summary_table.rows[i].cells[1].text = val
    if key and not val:
        # 分类标题行
        set_cell_shading(summary_table.rows[i].cells[0], 'D9E2F3')
        set_cell_shading(summary_table.rows[i].cells[1], 'D9E2F3')
        for cell in summary_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
    elif i == 0:
        set_cell_shading(summary_table.rows[i].cells[0], '4472C4')
        set_cell_shading(summary_table.rows[i].cells[1], '4472C4')
        for cell in summary_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.bold = True

set_table_font(summary_table)

# 保存文档
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(BASE_DIR, '作业分析报告.docx')
doc.save(output_path)
print(f'报告已生成: {output_path}')
